from docx import Document
import pandas as pd


def export_docx(
    df: pd.DataFrame,
    deals: pd.DataFrame,
    path: str = "market_report.docx",
) -> None:
    """
    Export real estate market report to DOCX.
    """
    doc = Document()

    # =====================================================
    # TITLE
    # =====================================================
    doc.add_heading("Hà Nội – Báo cáo thị trường bất động sản", level=1)

    # =====================================================
    # MARKET OVERVIEW
    # =====================================================
    avg_price_billion = df["price"].mean() / 1_000_000_000
    avg_price_m2 = df["price_million_per_m2"].mean()

    doc.add_heading("Tổng quan thị trường", level=2)
    doc.add_paragraph(
        f"Tổng số tin: {len(df)}\n"
        f"Giá trung bình: {avg_price_billion:.2f} tỷ\n"
        f"Giá trung bình / m²: {avg_price_m2:.1f} triệu/m²"
    )

    # =====================================================
    # GOOD DEALS
    # =====================================================

    number_of_hot_deals = 5 
    doc.add_heading("Tin giá tốt", level=2)

    doc.add_heading("Phương pháp xác định", level=3)
    doc.add_paragraph(
        """
- So sánh giá trên mỗi mét vuông (giá/m²) của từng tin với giá/m² trung vị của quận/huyện tương ứng.
- Chỉ giữ lại các tin có giá thấp hơn 25% so với mặt bằng chung của khu vực.
        """
    )

    doc.add_heading("Danh sách bất động sản giá tốt", level=3)


    top_deals = deals.head(number_of_hot_deals)

    for _, row in top_deals.iterrows():
        price_billion = row["price"] / 1_000_000_000

        doc.add_paragraph(
            f"- {row['title']} | {row['area_name']} | "
            f"{price_billion:.2f} tỷ | "
            f"{row['price_million_per_m2']:.1f} triệu/m²"
        )

    # =====================================================
    # SAVE
    # =====================================================
    doc.save(path)
    print(f"📄 Report saved to {path}")
